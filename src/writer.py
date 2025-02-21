import os
import sys
import zipfile
import io
import csv
from lxml import etree
from docx import Document
from PIL import Image


class DocumentAnalyzer:
    """
    A class for extracting and analyzing various components of a DOCX document.

    This class provides methods to extract information such as:
      - Paragraph details: text, length, style, font formatting (bold, italic, underline),
        font names, font sizes, and paragraph alignment.
      - Image details: image format and dimensions.
      - Table details: number of rows and columns.
      - Margin settings by parsing the document's XML structure.
    
    Libraries used:
      - python-docx for reading and interacting with DOCX files.
      - lxml for parsing XML content from DOCX.
      - Pillow for image processing.
    """

    def __init__(self, file_path: str):
        """
        Initialize the DocumentAnalyzer with a DOCX file.

        :param file_path: Path to the DOCX file to be analyzed.
        """
        self.doc = Document(file_path)
        self.docx_path = file_path

    def get_paragraph_alignment(self, paragraph) -> str:
        """
        Retrieve the alignment of a paragraph.

        This method maps the numerical alignment value from the DOCX paragraph
        to a human-readable string. If the attribute is not present or recognized,
        it returns "unknown".

        :param paragraph: A paragraph object from the DOCX document.
        :return: A string indicating the paragraph alignment 
                 ("left", "center", "right", "justified", or "unknown").
        """
        alignment_map = {0: "left", 1: "center", 2: "right", 3: "justified"}
        try:
            alignment = paragraph.alignment
            return alignment_map.get(alignment, "unknown")
        except AttributeError:
            # If the paragraph object does not have an alignment attribute.
            return "unknown"

    def get_paragraphs_info(self) -> tuple[list[dict], int]:
        """
        Extract detailed information from all paragraphs in the document.

        For each paragraph, this method extracts:
          - The stripped text (leading and trailing spaces removed)
          - The length of the text
          - The style name (e.g., Normal, Heading 1)
          - Boolean flags indicating whether the text is bold, italic, or underlined
          - A list of font names used in the runs of the paragraph
          - A list of font sizes (in points) used in the runs
          - The alignment of the paragraph

        Additionally, it counts how many paragraphs are empty (contain no text after stripping).

        :return: A tuple containing:
                 - A list of dictionaries with detailed information for each paragraph.
                 - An integer count of empty paragraphs.
        """
        paragraphs_data = []
        empty_lines_count = 0

        for paragraph in self.doc.paragraphs:
            # Check for empty paragraphs (after stripping whitespace).
            if not paragraph.text.strip():
                empty_lines_count += 1

            # Construct a dictionary with detailed paragraph info.
            paragraph_info = {
                "text": paragraph.text.strip(),
                "length": len(paragraph.text.strip()),
                "style": paragraph.style.name,
                "bold": any(run.bold for run in paragraph.runs),
                "italic": any(run.italic for run in paragraph.runs),
                "underline": any(run.underline for run in paragraph.runs),
                "font": [run.font.name for run in paragraph.runs if run.font.name],
                "size": [run.font.size.pt for run in paragraph.runs if run.font.size],
                "alignment": self.get_paragraph_alignment(paragraph)
            }
            paragraphs_data.append(paragraph_info)

        return paragraphs_data, empty_lines_count

    def get_images_info(self) -> list[dict]:
        """
        Extract details of all images embedded in the DOCX document.

        The method iterates through the document's relationships to find image files.
        It then uses Pillow to open each image from a byte stream, extracting the image
        format and its dimensions (width and height).

        :return: A list of dictionaries, each containing:
                 - "format": The image format (e.g., JPEG, PNG).
                 - "dimensions": A tuple with the image's width and height.
        """
        images_data = []
        for rel in self.doc.part.rels:
            # Check if the relationship refers to an image.
            if "image" in self.doc.part.rels[rel].target_ref:
                image_part = self.doc.part.rels[rel].target_part
                image_stream = io.BytesIO(image_part.blob)
                with Image.open(image_stream) as img:
                    images_data.append({
                        "format": img.format,
                        "dimensions": img.size
                    })
        return images_data

    def get_tables_info(self) -> list[dict]:
        """
        Extract information about all tables in the document.

        For every table, this method calculates the number of rows and columns.

        :return: A list of dictionaries, each containing:
                 - "rows": The count of rows in the table.
                 - "columns": The count of columns in the table.
        """
        tables_data = []
        for table in self.doc.tables:
            num_rows = len(table.rows)
            num_columns = len(table.columns)
            tables_data.append({
                "rows": num_rows,
                "columns": num_columns
            })
        return tables_data

    def get_margins(self) -> dict:
        """
        Extract the margin settings from the DOCX document.

        This method opens the DOCX file as a ZIP archive and reads the "word/document.xml"
        file, which contains the document's XML structure. It then parses the XML using lxml,
        locates the "pgMar" element (which holds margin values), and constructs a dictionary
        of these margin attributes.

        :return: A dictionary of margin settings (e.g., top, bottom, left, right) as defined in the XML.
                 If no margins are found, returns an empty dictionary.
        """
        margins_data = {}
        with zipfile.ZipFile(self.docx_path, "r") as docx_zip:
            with docx_zip.open("word/document.xml") as xml_file:
                xml_tree = etree.parse(xml_file)
                # Define the XML namespace used in DOCX documents.
                namespaces = {"w": "http://schemas.openxmlformats.org/wordprocessingml/2006/main"}
                # Extract the margin settings element.
                page_margins = xml_tree.xpath("//w:sectPr/w:pgMar", namespaces=namespaces)
                if page_margins:
                    margins_data = {attr: page_margins[0].get(attr) for attr in page_margins[0].keys()}
        return margins_data


class DocumentComparer:
    """
    A class for comparing two DOCX documents based on their structural and formatting elements.

    The class uses DocumentAnalyzer to extract features from both a reference document and a test document.
    It then compares these features—such as paragraphs, images, tables, and margins—and calculates a
    match percentage for each section. Finally, it generates a comprehensive Markdown report summarizing the differences.
    """

    def __init__(self, reference_file: str, test_file: str):
        """
        Initialize the DocumentComparer with a reference and a test document.

        :param reference_file: Path to the reference DOCX file (e.g., the solution file).
        :param test_file: Path to the DOCX file to be compared (e.g., a student's submission).
        """
        self.reference_analyzer = DocumentAnalyzer(reference_file)
        self.test_analyzer = DocumentAnalyzer(test_file)

    def assign_score(self, correct: int, total: int) -> float:
        """
        Calculate the score for a section based on the ratio of matching elements.

        If no elements are present for comparison (i.e., total is zero), the method returns 100%.

        :param correct: Number of matching elements.
        :param total: Total number of elements compared.
        :return: A percentage score (0 to 100).
        """
        if total == 0:
            return 100.0
        return (correct / total) * 100.0

    def compare_elements(self, reference_elements: list, test_elements: list,
                         element_name: str) -> tuple[list[str], float]:
        """
        Compare two lists of elements (e.g., paragraphs, images, tables, margins) from the documents.

        This method compares corresponding elements from the reference and test documents.
        It builds a list of differences for any mismatches and calculates a match percentage.
        Special conditions include:
          - Both lists empty: returns a perfect 100% match.
          - One list empty: returns 0% match.

        :param reference_elements: List of elements from the reference document.
        :param test_elements: List of elements from the test document.
        :param element_name: A string describing the element type (used in reporting differences).
        :return: A tuple containing:
                 - A list of strings describing differences.
                 - A float representing the match percentage (0 to 100).
        """
        differences_list = []
        correct_matches = 0
        total_elements = len(reference_elements)

        # Handle cases where one or both element lists are empty.
        if total_elements == 0 and len(test_elements) == 0:
            return differences_list, 100.0  # Both absent: perfect match.
        if total_elements == 0 or len(test_elements) == 0:
            return differences_list, 0.0    # One is missing: no match.

        # Compare each element pair.
        for index, (ref_elem, test_elem) in enumerate(zip(reference_elements, test_elements)):
            if ref_elem == test_elem:
                correct_matches += 1
            else:
                differences_list.append(
                    f"Difference in {element_name} {index + 1}: {ref_elem} ≠ {test_elem}"
                )

        score_percentage = self.assign_score(correct_matches, total_elements)
        return differences_list, score_percentage

    def tolerance_empty_lines(self, reference_empty: int, test_empty: int, tolerance: int = 1) -> bool:
        """
        Determine if the difference in the number of empty paragraphs (lines) is within an acceptable range.

        This method allows for minor discrepancies in empty lines between the documents,
        which can be acceptable for comparison purposes.

        :param reference_empty: Count of empty paragraphs in the reference document.
        :param test_empty: Count of empty paragraphs in the test document.
        :param tolerance: The maximum allowed difference (default is 1).
        :return: True if the difference is within the tolerance; False otherwise.
        """
        return abs(reference_empty - test_empty) <= tolerance

    def generate_markdown_report(self, report_lines: list[str]) -> str:
        """
        Generate a Markdown formatted report from a list of report lines.

        Each line is prefixed with a dash, and the report begins with a header.

        :param report_lines: List of report strings.
        :return: A single string containing the full Markdown report.
        """
        markdown_report = "# Comparison Report\n"
        for line in report_lines:
            markdown_report += f"- {line}\n"
        return markdown_report

    def compare_documents(self) -> tuple[str, float]:
        """
        Compare the reference and test DOCX documents, section by section, and produce a Markdown report.

        The comparison includes:
          - Paragraphs: text, formatting, and structure.
          - Images: formats and dimensions.
          - Tables: row and column counts.
          - Margins: document layout settings extracted from XML.
        
        For paragraphs, a bonus is added to the score if the empty line count difference is within tolerance.
        The final score is computed as the average of the individual section scores, capped at 100%.

        :return: A tuple containing:
                 - A Markdown formatted string summarizing the comparison results.
                 - The final score (float) representing overall match percentage.
        """
        report_lines = []

        # --- Compare Paragraphs ---
        ref_paragraphs, ref_empty = self.reference_analyzer.get_paragraphs_info()
        test_paragraphs, test_empty = self.test_analyzer.get_paragraphs_info()
        paragraph_differences, paragraph_score = self.compare_elements(
            ref_paragraphs, test_paragraphs, "paragraph"
        )
        # Apply a bonus if the empty line counts are within acceptable tolerance.
        if self.tolerance_empty_lines(ref_empty, test_empty, tolerance=1):
            paragraph_score = min(paragraph_score + 10, 100)  # Ensure the score does not exceed 100.
        report_lines.append(f"Paragraphs: {paragraph_score}% match")
        report_lines.extend(paragraph_differences)

        # --- Compare Images ---
        image_differences, image_score = self.compare_elements(
            self.reference_analyzer.get_images_info(),
            self.test_analyzer.get_images_info(),
            "image"
        )
        report_lines.append(f"Images: {image_score}% match")
        report_lines.extend(image_differences)

        # --- Compare Tables ---
        table_differences, table_score = self.compare_elements(
            self.reference_analyzer.get_tables_info(),
            self.test_analyzer.get_tables_info(),
            "table"
        )
        report_lines.append(f"Tables: {table_score}% match")
        report_lines.extend(table_differences)

        # --- Compare Margins ---
        # Wrap the margin dictionaries in lists for a consistent comparison format.
        margin_differences, margin_score = self.compare_elements(
            [self.reference_analyzer.get_margins()],
            [self.test_analyzer.get_margins()],
            "margins"
        )
        report_lines.append(f"Margins: {margin_score}% match")
        report_lines.extend(margin_differences)

        # --- Compute Final Score ---
        final_score = min((paragraph_score + image_score + table_score + margin_score) / 4, 100)
        report_lines.append(f"\nFinal Score: {final_score}%")

        # Generate and return the final Markdown report along with the final score.
        return self.generate_markdown_report(report_lines), final_score


class DocxAssignmentEvaluator:
    """
    A class to manage the evaluation of student DOCX submissions for a given assignment.

    The evaluator uses an assignment_identifier to locate the proper subfolder in both the
    assignments and solutions directories. It then compares each student's DOCX file (in the
    assignments folder) against the reference solution DOCX (in the corresponding solutions folder)
    using DocumentComparer. Finally, it generates a consolidated CSV report with each student's score.
    """

    def __init__(self, assignment_identifier: str):
        """
        Initialize the evaluator with the given assignment identifier.

        :param assignment_identifier: Unique identifier for the assignment; this is used
                                      to select the proper subfolders in both the assignments
                                      and solutions directories.
        """
        self.assignment_id = assignment_identifier
        self.assignments_dir = "assignments"
        self.solutions_dir = "solutions"
        self.evaluations_dir = "evaluations"

        # Construct the folder paths for the student's submissions and the solution.
        self.assignment_folder = os.path.join(self.assignments_dir, self.assignment_id)
        self.solution_file = os.path.join(self.solutions_dir, self.assignment_id, "solution.docx")
        # The report file will be generated in the evaluations directory.
        self.report_file = os.path.join(self.evaluations_dir, f"{self.assignment_id}-Evaluation_Report.csv")

    def verify_resources(self) -> bool:
        """
        Verify that the necessary folders and files exist for the evaluation process.

        Checks for:
          - Existence of the assignment folder (student submissions).
          - Existence of the solution DOCX file in the solutions directory.
          - Existence (or creatability) of the evaluations directory.

        :return: True if all required resources exist; False otherwise.
        """
        if not os.path.exists(self.assignment_folder):
            print(f"Error: Assignment folder '{self.assignment_folder}' not available.")
            return False

        if not os.path.exists(self.solution_file):
            print(f"Error: Solution file '{self.solution_file}' not available.")
            return False

        # Ensure evaluations directory exists (or can be created).
        os.makedirs(self.evaluations_dir, exist_ok=True)
        return True

    @staticmethod
    def extract_student_name(file_path: str) -> tuple[str, str]:
        """
        Extract the student's first name and surname from the DOCX file name.

        Assumes the file name is in the format "firstname-surname.docx".

        :param file_path: The path to the student's DOCX file.
        :return: A tuple containing the first name and surname (both capitalized).
        """
        base_name = os.path.splitext(os.path.basename(file_path))[0]
        name_parts = base_name.split('-')
        if len(name_parts) >= 2:
            first_name, surname = [part.lower().capitalize() for part in name_parts[:2]]
        else:
            # If the file name doesn't follow the expected format, use the whole name as first name.
            first_name, surname = base_name.lower().capitalize(), ""
        return first_name, surname

    def run_evaluation(self) -> None:
        """
        Execute the evaluation process for all student DOCX submissions in the assignment folder.

        For each DOCX file in the assignment folder, the method compares it against the reference solution
        DOCX using DocumentComparer. The final score is recorded, and a consolidated CSV report is generated.
        """
        if not self.verify_resources():
            return

        evaluation_results = []

        # Iterate over each file in the assignment folder.
        for file in os.listdir(self.assignment_folder):
            if file.endswith(".docx"):
                student_file_path = os.path.join(self.assignment_folder, file)
                # Use DocumentComparer to compare the student's submission with the solution.
                comparer = DocumentComparer(self.solution_file, student_file_path)
                report, final_score = comparer.compare_documents()

                # Extract student's name from the file name.
                first_name, surname = self.extract_student_name(student_file_path)
                evaluation_results.append({
                    "Name": first_name,
                    "Surname": surname,
                    "Score (%)": final_score
                })

                # Optionally, save the detailed markdown report for each student.
                individual_report_path = os.path.join(
                    self.evaluations_dir, f"{self.assignment_id}_{first_name}_{surname}_Report.md"
                )
                with open(individual_report_path, "w") as report_file:
                    report_file.write(report)

                print(f"Evaluated {file}: {final_score}% (detailed report at {individual_report_path})")

        # Write a consolidated CSV report.
        with open(self.report_file, "w", newline="") as csvfile:
            fieldnames = ["Name", "Surname", "Score (%)"]
            writer = csv.DictWriter(csvfile, fieldnames=fieldnames)
            writer.writeheader()
            for result in evaluation_results:
                writer.writerow(result)

        print(f"Consolidated Evaluation Report generated at: {self.report_file}")


def main():
    """
    Main function to run the DOCX assignment evaluation.

    The assignment identifier is taken from the command-line arguments (if provided).
    This identifier is used to select the appropriate subfolders from the assignments
    and solutions directories.
    """
    assignment_identifier = sys.argv[1] if len(sys.argv) > 1 else None
    if not assignment_identifier:
        print("Error: Assignment identifier must be provided as an input parameter.")
        return

    evaluator = DocxAssignmentEvaluator(assignment_identifier)
    evaluator.run_evaluation()


if __name__ == "__main__":
    main()
